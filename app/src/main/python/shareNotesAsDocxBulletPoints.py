import io
from docx import Document
from docx.shared import Pt

def export_bullet_points_to_docx(title, bullet_points):
    # Create a new Document
    doc = Document()

    # Set document title
    doc.add_heading(title, level=1)

    # Add bullet points
    for point in bullet_points:
        # Handle long/multiline text by splitting it into lines
        lines = point.split('\n')
        for line in lines:
            # Add each line as a bullet point
            paragraph = doc.add_paragraph(line, style='ListBullet')
            # Optionally set font size if needed
            for run in paragraph.runs:
                run.font.size = Pt(12)

    # Write the document to an in-memory bytes buffer
    byte_stream = io.BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)  # Move the cursor to the beginning of the stream

    return byte_stream.read()  # Return the byte content of the file
