import io
from docx import Document
from docx.shared import Pt

def export_to_docx_in_memory(title, checkbox_states, checkbox_texts):
    # Create a new Document
    doc = Document()

    # Set document title
    doc.add_heading(title, 0)

    # Add checkboxes and text
    for i, (state, text) in enumerate(zip(checkbox_states, checkbox_texts)):
        # Use text representation for checkboxes
        checkbox_symbol = '[x]' if state else '[ ]'  # Checked and unchecked text representations

        # Add checkbox and text
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"{checkbox_symbol} ")
        run.font.size = Pt(12)

        # Handle long/multine text by splitting it into lines
        lines = text.split('\n')
        for line in lines:
            paragraph.add_run(line).font.size = Pt(12)
            paragraph.add_run('\n')  # Add a new line for multiline support

    # Write the document to an in-memory bytes buffer
    byte_stream = io.BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)  # Move the cursor to the beginning of the stream

    return byte_stream.read()  # Return the byte content of the file
